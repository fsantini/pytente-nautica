import json
import os
import docx

def main():
    # Open the docx file
    doc = docx.Document('allegato-a-dd-106-del-120522.docx')
    #doc = docx.Document('elenco-quiz-vela.docx')

    print(doc.tables)

    # Find the first table in the document
    table = doc.tables[0]

    # Create a list to store the content of each cell in the table
    table_data = []

    os.makedirs('images', exist_ok=True)

    images = {}
    for inline_shape in doc.inline_shapes:
        try:
            key = inline_shape._inline.attrib['{http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing}anchorId']
        except KeyError:
            continue
        rId = inline_shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        image = doc.part.related_parts[rId]._blob
        images[rId] = image

    for key, part in doc.part.related_parts.items():
        images[key] = part._blob

    print(list(images.keys()))

    do_print = False

    for table_index, table in enumerate(doc.tables):
        # Loop over the rows in the table
        for i, row in enumerate(table.rows):
            row_data = {}
            row_data['answers'] = []
            row_data['right_answer'] = None
            skip_row = False
            # Loop over the cells in the row
            for j, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if j == 0:
                    # this should be the question number
                    try:
                        q_num = int(txt)
                        row_data['num'] = q_num
                        if q_num == 130:
                            print(table_index,i)
                    except ValueError:
                        # this is not a question
                        skip_row = True
                        break
                elif j == 1: # might contain an image
                    id = None
                    if 'w:drawing' in cell._tc.xml:
                        blipFill = cell._tc.xpath('.//w:drawing/wp:inline/a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip')[0]
                        id = blipFill.attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed']
                    if 'w:pict' in cell._tc.xml:
                        #anchorId = cell._tc.xpath('.//w:pict')[0].attrib['{http://schemas.microsoft.com/office/word/2010/wordml}anchorId']
                        id =row.cells[1]._tc.xpath('.//w:pict')[0].find('{urn:schemas-microsoft-com:vml}shape').find('{urn:schemas-microsoft-com:vml}imagedata').attrib['{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id']
                    if id:
                        image = images[id]
                        filename = f'{table_index}_{i}_{j}.png'
                        with open(os.path.join('images', filename), 'wb') as f:
                            f.write(image)
                        row_data['image'] = filename
                    else:
                        row_data['image'] = None
                elif j == 2:
                    row_data['question'] = txt
                elif j == 3 or j == 5 or j == 7:
                    row_data['answers'].append(txt)
                elif j == 4 or j == 6 or j == 8:
                    if txt.upper() == 'V':
                        row_data['right_answer'] = int((j - 4)/2)
            # Append the row data list to the table data list
            if skip_row:
                continue
            print(row_data)
            table_data.append(row_data)

    with open('questions.json', 'w') as f:
        json.dump(table_data, f, indent=4)

if __name__ == '__main__':
    main()
